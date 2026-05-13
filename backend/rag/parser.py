"""
=============================================================================
RAG 核心模块 — 多格式文档解析器 (Document Parser)
=============================================================================

【RAG 概念入门】
这是 Ingest 管线的第一步：把用户上传的各种格式文件转成纯文本。
不同格式的"壳子"不一样，但内核都是文字。Parser 负责剥壳取肉。

【类比 Java 的接口设计】
这里的 DocumentParser 使用了策略模式(Strategy Pattern)的变体：
- 统一入口 parse() 根据文件类型分发到不同的解析器
- 每种格式有自己的 _parse_xxx() 方法
- 返回统一的数据结构 ParsedDocument

在 Java 中会这样设计：
    interface DocumentParser {
        ParsedDocument parse(byte[] content, String filename);
    }
    class PdfParser implements DocumentParser { ... }
    class DocxParser implements DocumentParser { ... }

这里用 Python 的 classmethod + 字典映射简化了实现。
=============================================================================
"""

from typing import Optional
from dataclasses import dataclass, field
from io import BytesIO


@dataclass
class ParsedDocument:
    """
    解析后的文档数据结构。

    Python 的 @dataclass 等价于 Java 的 Lombok @Data 或 Kotlin 的 data class：
    自动生成 __init__、__repr__、__eq__ 等方法。

    字段说明：
    - text:     提取出的纯文本正文
    - title:    文档标题（从文件名或元数据推测）
    - metadata: 附加元数据，如 PDF 的作者、创建时间等
    """
    text: str
    title: Optional[str] = None
    # field(default_factory=dict) 保证每个实例有自己的空字典，
    # 避免 Python 中可变的默认参数被所有实例共享的坑。
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    """
    多格式文档解析器。

    【支持的文件类型映射表】
    键 = MIME type（HTTP 上传时的 Content-Type）
    值 = 文件后缀

    类比 Java 的 Map<String, String>
    """
    SUPPORTED_TYPES = {
        "application/pdf": ".pdf",
        # .docx 的 MIME type 就是这么长，Microsoft 定义的
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/markdown": ".md",
        "text/html": ".html",
        "text/plain": ".txt",
    }

    # =========================================================================
    # 统一入口：类似 Java 的工厂方法 (Factory Method)
    # =========================================================================

    @classmethod
    def parse(cls, content: bytes, filename: str, mime_type: str = "") -> ParsedDocument:
        """
        根据文件名/MIME 类型自动选择解析器，统一入口。

        【为什么用 @classmethod】
        因为 DocumentParser 没有实例状态（不需要加载模型），
        所有方法都是纯逻辑操作。所以直接用类方法，不需要 new DocumentParser()。
        类似于 Java 的 static 方法。

        Args:
            content:   文件的原始字节内容（从 UploadFile.read() 获得）
            filename:  原始文件名，如 "技术文档.pdf"
            mime_type: HTTP 的 Content-Type，如 "application/pdf"

        Returns:
            ParsedDocument: 统一格式的解析结果
        """
        ext = cls._get_ext(filename, mime_type)

        # 策略分发 —— 每个 elif 分支就是一个"策略"
        if ext == ".pdf":
            return cls._parse_pdf(content, filename)
        elif ext == ".docx":
            return cls._parse_docx(content, filename)
        elif ext == ".md":
            return cls._parse_markdown(content, filename)
        elif ext == ".html":
            return cls._parse_html(content, filename)
        else:
            # 默认按纯文本处理（兜底策略）
            return cls._parse_text(content, filename)

    @classmethod
    def _get_ext(cls, filename: str, mime_type: str) -> str:
        """
        推断文件类型。优先 MIME type（更可靠），其次看文件后缀。
        比如 "application/pdf" 比 ".pdf" 后缀更可信（后缀可以被伪造）。
        """
        if mime_type in cls.SUPPORTED_TYPES:
            return cls.SUPPORTED_TYPES[mime_type]
        # 后备方案：检查文件后缀
        filename_lower = filename.lower()
        for ext in [".pdf", ".docx", ".md", ".html", ".txt"]:
            if filename_lower.endswith(ext):
                return ext
        return ".txt"  # 最终兜底

    # =========================================================================
    # 各格式解析器
    # =========================================================================

    @classmethod
    def _parse_pdf(cls, content: bytes, filename: str) -> ParsedDocument:
        """
        PDF 解析。

        【PDF 格式的特点】
        PDF 不是纯文本格式，它是"页面描述语言"——定位每个字符在纸上的坐标。
        所以我们不能直接读字节，需要通过 pypdf 库来提取。

        类比：PDF 像是一张"打印出来的纸"的照片——你能看到字，
        但要取出来需要 OCR 或专门的提取工具。pypdf 就是那个工具。

        【为什么用 BytesIO】
        BytesIO 把 bytes 包装成"虚拟文件"，让 pypdf 以为在读取磁盘文件。
        类似 Java 的 ByteArrayInputStream。
        """
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        parts: list[str] = []
        metadata: dict = {}

        # PDF 元数据：作者、标题、创建时间等
        if reader.metadata:
            title = reader.metadata.get("/Title", None)
            if title:
                metadata["title"] = title

        # 逐页提取文字
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)

        # 用两个换行符连接各页，保持页面间的分隔
        return ParsedDocument(
            text="\n\n".join(parts),
            title=metadata.get("title") or filename,
            metadata=metadata,
        )

    @classmethod
    def _parse_docx(cls, content: bytes, filename: str) -> ParsedDocument:
        """
        Word (.docx) 解析。

        【.docx 格式的本质】
        .docx 本质上是一个 ZIP 压缩包，里面是 XML 文件。
        python-docx 帮我们解压并解析这些 XML。
        类比：一个 .docx 文件 = 把多个 config.xml 打包成 zip。
        """
        from docx import Document

        doc = Document(BytesIO(content))
        parts: list[str] = []

        # 遍历所有段落，跳过空段落
        # 每个 paragraph 是文档中的一个自然段
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        return ParsedDocument(
            text="\n".join(parts),
            title=filename,
        )

    @classmethod
    def _parse_markdown(cls, content: bytes, filename: str) -> ParsedDocument:
        """
        Markdown 解析。

        Markdown 本身是纯文本，不需要特殊解析。
        但我们尝试提取一级标题（# 开头的行）作为文档标题。
        注意区分 # 一级标题 和 ## 二级标题。
        """
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")
        title = filename

        # 找第一个一级标题（只匹配 # 不匹配 ##）
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# ") and not stripped.startswith("## "):
                title = stripped[2:].strip()
                break

        return ParsedDocument(text=text, title=title)

    @classmethod
    def _parse_html(cls, content: bytes, filename: str) -> ParsedDocument:
        """
        HTML 解析。

        BeautifulSoup 解析 HTML DOM 树，提取可见文本。
        类比 Java 中的 Jsoup：同样用于解析 HTML 提取文本。

        【get_text 的 separator 参数】
        指定不同元素间的分隔符为换行，避免文字粘连。
        比如 <p>第一段</p><p>第二段</p> 会变成 "第一段\n第二段"
        而不是 "第一段第二段"。
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, "html.parser")

        # 尝试从 <title> 标签获取标题
        title_tag = soup.find("title")
        title = title_tag.get_text(strip=True) if title_tag else filename

        return ParsedDocument(
            text=soup.get_text(separator="\n", strip=True),
            title=title,
        )

    @classmethod
    def _parse_text(cls, content: bytes, filename: str) -> ParsedDocument:
        """
        纯文本解析（兜底）。
        尝试 UTF-8 解码，失败则用 replace 策略跳过乱码字符。
        """
        text = content.decode("utf-8", errors="replace")
        return ParsedDocument(text=text, title=filename)